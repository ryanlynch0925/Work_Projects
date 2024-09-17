import pandas as pd
import os
import win32com.client as win32
from docx import Document

successfully_generated_pdfs = []

# Read the Excel file and filter the data
excel_file = r"C:\Users\DavidLynch\OneDrive - Tidal Wave Autospa\Documents\Utitlies Project\3.xlsx"
df = pd.read_excel(excel_file, sheet_name='Print', engine='openpyxl')

df['Site Name'] = df['Site Name'].str.replace('/', '_')
filtered_df = df[(df['Created?'] == 'No') & (df['Ready?'] == 'Ready')]
# print(df.head())
# print(filtered_df.head())

# Specify the path to save the PDF files and the Word document for errors
output_folder = r"C:\Users\DavidLynch\OneDrive - Tidal Wave Autospa\Documents\Utitlies Project\PDFs"
error_doc_path = r"C:\Users\DavidLynch\OneDrive - Tidal Wave Autospa\Documents\Utitlies Project\Errors.docx"

if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Create or open a Word document for errors
error_doc = Document()
error_doc.add_heading("Errors in PDF Generation", level=1)

# Create an Excel application
excel = win32.gencache.EnsureDispatch('Excel.Application')
excel.Visible = False

# Load the Excel file
workbook = excel.Workbooks.Open(excel_file)

# Iterate through filtered data
for index, row in filtered_df.iterrows():
    site_name = row['Site Name']
    try:
        if not pd.isna(site_name) and site_name:  # Check if the site name is not empty
            # Write the site name to cell B11 in 'form' sheet
            workbook.Sheets('Form').Cells(11, 2).Value = site_name  # Assumes B11 is cell (11, 2)
            # Save the 'Form' sheet as a PDF
            pdf_file_path = os.path.join(output_folder, f"{site_name}.pdf")
            workbook.Sheets('Form').ExportAsFixedFormat(0, pdf_file_path)  # 0 indicates PDF format

            successfully_generated_pdfs.append(f"{site_name}.pdf")
    except Exception as e:
        # Handle the error and add the site name to the Word document for errors
        error_message = f"Error for site '{site_name}': {str(e)}"
        error_doc.add_paragraph(error_message)

# Save the Word document for errors
error_doc.save(error_doc_path)

# Close the Excel application
workbook.Close(SaveChanges=False)
excel.Quit()

# Create a new section in the Word document and add a list of successfully generated PDFs
success_doc = Document(error_doc_path)
success_doc.add_section()
success_doc.add_heading("Successfully Generated PDFs", level=1)

for pdf_name in successfully_generated_pdfs:
    success_doc.add_paragraph(pdf_name)

# Save the updated Word document
success_doc.save(error_doc_path)


